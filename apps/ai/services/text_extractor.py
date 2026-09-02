"""Resume text extraction from PDF, DOCX, and image files."""

import re
import io
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {str(e)}")


def extract_text_from_image(file_bytes: bytes) -> str:
    """Extract text from image using OCR."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(img)
        return text
    except ImportError:
        raise ValueError("OCR not available — pytesseract not installed")
    except Exception as e:
        raise ValueError(f"Image OCR failed: {str(e)}")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Route to appropriate extractor based on file type."""
    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "txt": lambda b: b.decode("utf-8", errors="replace"),
        "image": extract_text_from_image,
    }
    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")
    text = extractor(file_bytes)
    if not text or len(text.strip()) < 10:
        raise ValueError(f"Insufficient text extracted from {file_type} file")
    return text.strip()
